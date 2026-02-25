"""
DOCX Exporter

Exports:
- DataFrame to Word table
- Dictionary (metrics/report) to Word
- Matplotlib figure to Word
"""

from __future__ import annotations

from typing import Dict, Any, Optional
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

import matplotlib.figure

from ..context import Context


class DOCXExporter:
    """
    Handles exporting data and reports to Word (.docx).
    """

    def __init__(self, context: Context) -> None:
        self.context = context

    # -------------------------------------------------
    # Export DataFrame
    # -------------------------------------------------

    def toDOCX(
        self,
        path: str,
        title: Optional[str] = None,
        max_rows: int = 50,
    ) -> "DOCXExporter":
        """
        Export current DataFrame to Word table.
        """
        self.context.ensure_dataframe()
        df: pd.DataFrame = self.context.dataframe

        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()

        if title:
            document.add_heading(title, level=1)

        df_subset = df.head(max_rows)

        table = document.add_table(
            rows=len(df_subset) + 1,
            cols=len(df_subset.columns),
        )

        # Header
        for col_idx, col_name in enumerate(df_subset.columns):
            table.rows[0].cells[col_idx].text = str(col_name)

        # Data
        for row_idx, row in enumerate(df_subset.values):
            for col_idx, cell_value in enumerate(row):
                table.rows[row_idx + 1].cells[col_idx].text = str(cell_value)

        document.save(output_path)

        return self

    # -------------------------------------------------
    # Export Dictionary / Metrics
    # -------------------------------------------------

    def exportDict(
        self,
        data: Dict[str, Any],
        path: str,
        title: str = "Report",
    ) -> "DOCXExporter":
        """
        Export dictionary content into Word document.
        """
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(title, level=1)

        for key, value in data.items():
            document.add_paragraph(f"{key}: {value}")

        document.save(output_path)

        return self

    # -------------------------------------------------
    # Export Matplotlib Figure
    # -------------------------------------------------

    def exportFigure(
        self,
        figure: matplotlib.figure.Figure,
        path: str,
        width: float = 6.0,
    ) -> "DOCXExporter":
        """
        Export matplotlib figure into Word document.
        """
        output_path = Path(path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        temp_image = output_path.with_suffix(".png")
        figure.savefig(temp_image)

        document = Document()
        document.add_picture(str(temp_image), width=Inches(width))
        document.save(output_path)

        temp_image.unlink(missing_ok=True)

        return self

    # -------------------------------------------------
    # Representation
    # -------------------------------------------------

    def __repr__(self) -> str:
        rows, cols = (0, 0)
        if self.context.dataframe is not None:
            rows, cols = self.context.dataframe.shape
        return f"<DOCXExporter rows={rows} cols={cols}>"
